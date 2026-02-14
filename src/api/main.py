"""
═══════════════════════════════════════════════════════════════════
  B2B Intelligence Platform — API Server
  Интеллектуальная собственность АО «Арагант Групп»
  Copyright (c) 2024-2026 АО «Арагант Групп». Все права защищены.
═══════════════════════════════════════════════════════════════════

FastAPI Backend for B2B Intelligence Platform.
Serves the API for the Mission Control dashboard.
"""
import os
import json
from typing import Optional
from datetime import datetime
from threading import Thread

from fastapi import FastAPI, Depends, UploadFile, File, HTTPException, Query, BackgroundTasks
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from sqlalchemy.orm import Session
from sqlalchemy import func, desc

from src.database import get_db
from src.database.models import Company, Person, Contact, Intelligence, Interaction, Document

app = FastAPI(
    title="B2B Intelligence Platform",
    description="API для управления разведкой и продажами B2B",
    version="0.2.0"
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Serve frontend static files
FRONTEND_DIR = os.path.join(os.path.dirname(__file__), '..', '..', 'frontend')
if os.path.isdir(FRONTEND_DIR):
    app.mount("/app", StaticFiles(directory=FRONTEND_DIR, html=True), name="frontend")


# ─── Company Profile (Our company context) ───
COMPANY_PROFILE_PATH = os.path.join(
    os.path.dirname(__file__), '..', 'data', 'company_profile.json'
)

@app.get("/api/v1/profile")
def get_company_profile():
    """Return our company's profile (for AI context)."""
    with open(COMPANY_PROFILE_PATH, 'r', encoding='utf-8') as f:
        return json.load(f)


# ─── Companies (Sellers / Leads) ───
@app.get("/api/v1/companies")
def list_companies(
    skip: int = 0, 
    limit: int = 50,
    status: Optional[str] = None,
    min_score: Optional[int] = None,
    search: Optional[str] = None,
    db: Session = Depends(get_db)
):
    """List companies with filtering and pagination."""
    q = db.query(Company)
    
    if status:
        q = q.filter(Company.enrichment_status == status)
    if min_score is not None:
        q = q.filter(Company.lead_score >= min_score)
    if search:
        q = q.filter(Company.name.ilike(f"%{search}%"))
    
    total = q.count()
    companies = q.order_by(desc(Company.lead_score)).offset(skip).limit(limit).all()
    
    return {
        "total": total,
        "items": [
            {
                "id": c.id,
                "key": c.key,
                "name": c.name,
                "legal_form": c.legal_form,
                "revenue_total": c.revenue_total,
                "sales_total": c.sales_total,
                "wb_present": c.wb_present,
                "ozon_present": c.ozon_present,
                "lead_score": c.lead_score,
                "enrichment_status": c.enrichment_status,
                "website": c.website,
                "contacts_count": len(c.contacts),
                "persons_count": len(c.persons),
            }
            for c in companies
        ]
    }


@app.get("/api/v1/companies/{company_id}")
def get_company_dossier(company_id: int, db: Session = Depends(get_db)):
    """Full dossier on a single company."""
    company = db.query(Company).filter(Company.id == company_id).first()
    if not company:
        raise HTTPException(status_code=404, detail="Company not found")
    
    return {
        "company": {
            "id": company.id,
            "key": company.key,
            "name": company.name,
            "legal_form": company.legal_form,
            "inn": company.inn,
            "revenue_total": company.revenue_total,
            "sales_total": company.sales_total,
            "avg_price": company.avg_price,
            "wb_present": company.wb_present,
            "ozon_present": company.ozon_present,
            "wb_brand_link": company.wb_brand_link,
            "ozon_brand_link": company.ozon_brand_link,
            "website": company.website,
            "lead_score": company.lead_score,
            "enrichment_status": company.enrichment_status,
        },
        "persons": [
            {"id": p.id, "name": p.full_name, "role": p.role, "source": p.source}
            for p in company.persons
        ],
        "contacts": [
            {"id": c.id, "type": c.type, "value": c.value, "label": c.label, "source": c.source, "verified": c.is_verified}
            for c in company.contacts
        ],
        "intelligence": {
            "brand_dna": company.intelligence.brand_dna if company.intelligence else None,
            "pain_points": company.intelligence.pain_points if company.intelligence else None,
            "approach_strategy": company.intelligence.approach_strategy if company.intelligence else None,
            "call_script": company.intelligence.call_script if company.intelligence else None,
            "score_breakdown": company.intelligence.score_breakdown if company.intelligence else None,
        } if company.intelligence else None,
        "interactions": [
            {"id": i.id, "type": i.type, "direction": i.direction, "status": i.status, 
             "summary": i.content_summary, "created_at": str(i.created_at)}
            for i in company.interactions
        ]
    }


@app.get("/api/v1/stats")
def get_dashboard_stats(db: Session = Depends(get_db)):
    """Dashboard summary statistics."""
    total = db.query(func.count(Company.id)).scalar()
    enriched = db.query(func.count(Company.id)).filter(Company.enrichment_status == 'enriched').scalar()
    hot_leads = db.query(func.count(Company.id)).filter(Company.lead_score >= 80).scalar()
    with_website = db.query(func.count(Company.id)).filter(Company.website.isnot(None)).scalar()
    
    return {
        "total_companies": total,
        "enriched": enriched,
        "hot_leads": hot_leads,
        "with_website": with_website,
        "enrichment_rate": round(enriched / total * 100, 1) if total else 0,
    }


# ─── Documents (КП, Спецификации, Договоры для RAG) ───
@app.post("/api/v1/documents/upload")
async def upload_document(
    file: UploadFile = File(...),
    doc_type: str = Query(default="other", description="kp_template, specification, contract, price_list, presentation, other"),
    db: Session = Depends(get_db)
):
    """Upload a document (PDF, DOCX, TXT) for the AI knowledge base."""
    content = await file.read()
    
    # Extract text based on file type
    text_content = ""
    filename = file.filename or "unknown"
    
    if filename.endswith('.txt'):
        text_content = content.decode('utf-8', errors='ignore')
    elif filename.endswith('.pdf'):
        try:
            import pdfplumber
            import io
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                text_content = "\n".join(page.extract_text() or "" for page in pdf.pages)
        except ImportError:
            text_content = "[PDF extraction requires pdfplumber]"
    elif filename.endswith('.docx'):
        try:
            import docx
            import io
            doc = docx.Document(io.BytesIO(content))
            text_content = "\n".join(p.text for p in doc.paragraphs)
        except ImportError:
            text_content = "[DOCX extraction requires python-docx]"
    else:
        text_content = content.decode('utf-8', errors='ignore')
    
    # Save to DB
    document = Document(
        filename=filename,
        doc_type=doc_type,
        content_text=text_content,
        doc_metadata={"size": len(content), "format": filename.split('.')[-1]}
    )
    db.add(document)
    db.commit()
    db.refresh(document)
    
    return {
        "id": document.id,
        "filename": document.filename,
        "doc_type": document.doc_type,
        "text_length": len(text_content),
        "status": "uploaded"
    }


@app.get("/api/v1/documents")
def list_documents(db: Session = Depends(get_db)):
    """List all uploaded documents."""
    docs = db.query(Document).order_by(desc(Document.uploaded_at)).all()
    return [
        {
            "id": d.id,
            "filename": d.filename,
            "doc_type": d.doc_type,
            "text_length": len(d.content_text) if d.content_text else 0,
            "uploaded_at": str(d.uploaded_at),
        }
        for d in docs
    ]


# ─── Enrichment ───
@app.post("/api/v1/enrich/{company_id}")
def enrich_single(company_id: int, use_ai: bool = False, db: Session = Depends(get_db)):
    """Enrich a single company (rule-based by default, add ?use_ai=true for GPT)."""
    from src.ai.brain import calculate_lead_score

    company = db.query(Company).get(company_id)
    if not company:
        raise HTTPException(status_code=404, detail="Company not found")

    company_dict = {
        "name": company.name,
        "revenue_total": company.revenue_total,
        "sales_total": company.sales_total,
        "wb_present": company.wb_present,
        "ozon_present": company.ozon_present,
        "avg_price": company.avg_price,
        "website": company.website,
    }

    score = calculate_lead_score(company_dict)
    company.lead_score = score
    company.enrichment_status = "enriched"
    db.commit()

    return {"id": company_id, "name": company.name, "lead_score": score, "status": "enriched"}


@app.post("/api/v1/enrich/batch")
def enrich_batch_endpoint(
    limit: int = Query(default=50),
    db: Session = Depends(get_db)
):
    """Batch enrich top companies by revenue (rule-based scoring)."""
    from src.ai.brain import calculate_lead_score

    companies = db.query(Company).filter_by(
        enrichment_status="new"
    ).order_by(Company.revenue_total.desc().nulls_last()).limit(limit).all()

    results = []
    for c in companies:
        company_dict = {
            "name": c.name,
            "revenue_total": c.revenue_total,
            "sales_total": c.sales_total,
            "wb_present": c.wb_present,
            "ozon_present": c.ozon_present,
            "avg_price": c.avg_price,
            "website": c.website,
        }
        score = calculate_lead_score(company_dict)
        c.lead_score = score
        c.enrichment_status = "enriched"
        results.append({"id": c.id, "name": c.name, "lead_score": score})

    db.commit()
    return {"enriched": len(results), "results": results}


# ─── Bitrix24 ───
@app.get("/api/v1/bitrix24/stats")
def bitrix24_stats():
    """Bitrix24 CRM quick stats."""
    try:
        from src.integrations.bitrix24 import Bitrix24Client
        client = Bitrix24Client()

        leads = client.call("crm.lead.list", {"select[]": "ID", "limit": 1})
        contacts = client.call("crm.contact.list", {"select[]": "ID", "limit": 1})
        deals = client.call("crm.deal.list", {"select[]": "ID", "limit": 1})
        companies = client.call("crm.company.list", {"select[]": "ID", "limit": 1})

        return {
            "connected": True,
            "leads": leads.get("total", 0),
            "contacts": contacts.get("total", 0),
            "deals": deals.get("total", 0),
            "companies": companies.get("total", 0),
        }
    except Exception as e:
        return {"connected": False, "error": str(e)}


@app.get("/api/v1/bitrix24/leads")
def bitrix24_leads(limit: int = 10):
    """Get recent leads from Bitrix24."""
    try:
        from src.integrations.bitrix24 import Bitrix24Client
        client = Bitrix24Client()
        leads = client.get_leads(limit=limit)
        return {"total": len(leads), "items": leads}
    except Exception as e:
        return {"error": str(e)}



# ─── CRM Analytics & Sync ───
@app.get("/api/v1/crm/analytics")
def get_crm_analytics(db: Session = Depends(get_db)):
    """Get latest CRM analytics report."""
    from sqlalchemy import text
    try:
        # Check if table exists first (in case sync hasn't run)
        db.execute(text("SELECT 1 FROM crm_analytics LIMIT 1"))
        
        result = db.execute(text("SELECT data FROM crm_analytics WHERE analysis_type='full_report' ORDER BY id DESC LIMIT 1")).fetchone()
        if not result:
            return {"error": "No analytics data found. Sync is running..."}
        return result[0]
    except Exception as e:
        return {"error": f"CRM data not available yet: {str(e)}"}


@app.get("/api/v1/crm/leads/lost")
def get_lost_leads(limit: int = 50, days: int = 30, db: Session = Depends(get_db)):
    """Get lost leads (no activity > N days)."""
    from sqlalchemy import text
    try:
        query = text("""
        SELECT id, title, name, last_name, company_title, phone, email, status_id, date_modify 
        FROM bitrix_leads 
        WHERE status_id NOT IN ('CONVERTED','JUNK') 
        AND date_modify < NOW() - INTERVAL :days || ' days'
        ORDER BY date_modify ASC 
        LIMIT :limit
        """)
        # Postgres INTERVAL syntax fix: '30 days' string concatenation
        
        result = db.execute(query, {"days": days, "limit": limit}).fetchall()
        
        # Manually map to dict if mappings() is not available or reliable across versions
        return {
            "items": [
                {
                    "id": row[0], "title": row[1], "name": row[2], "last_name": row[3],
                    "company_title": row[4], "phone": row[5], "email": row[6],
                    "status_id": row[7], "date_modify": str(row[8])
                }
                for row in result
            ]
        }
    except Exception as e:
        return {"error": str(e), "items": []}


@app.post("/api/v1/crm/sync")
def trigger_crm_sync(background_tasks: BackgroundTasks):
    """Trigger Bitrix24 sync in background."""
    def run_sync():
        import subprocess
        import sys
        # Use simple python command, assume env is set
        subprocess.run([sys.executable, "src/integrations/bitrix_sync.py"], capture_output=False)
    
    background_tasks.add_task(run_sync)
    return {"status": "Sync started in background"}

# ─── Recon — Web Parsing ───
@app.post("/api/v1/recon/crawl")
def trigger_recon_crawl(background_tasks: BackgroundTasks, limit: int = 50):
    """Запустить парсинг сайтов компаний в фоне."""
    def run_crawl():
        import subprocess
        import sys
        subprocess.run(
            [sys.executable, "scripts/recon_enrichment.py", "--limit", str(limit)],
            capture_output=False
        )
    
    background_tasks.add_task(run_crawl)
    return {"status": "Парсинг запущен в фоне", "limit": limit}


@app.get("/api/v1/recon/status")
def recon_status(db: Session = Depends(get_db)):
    """Статистика по парсингу сайтов."""
    from sqlalchemy import text
    try:
        total = db.execute(text("SELECT COUNT(*) FROM companies WHERE website IS NOT NULL AND website != ''")).scalar()
        crawled = db.execute(text("""
            SELECT COUNT(DISTINCT company_id) FROM contacts WHERE source = 'web_crawl'
        """)).scalar()
        total_contacts = db.execute(text("SELECT COUNT(*) FROM contacts WHERE source = 'web_crawl'")).scalar()
        
        return {
            "companies_with_website": total,
            "companies_crawled": crawled,
            "total_contacts_found": total_contacts,
            "remaining": total - crawled
        }
    except Exception as e:
        return {"error": str(e)}


@app.post("/api/v1/recon/crawl-one/{company_id}")
def crawl_single_company(company_id: int, db: Session = Depends(get_db)):
    """Спарсить сайт одной компании прямо сейчас."""
    company = db.query(Company).filter(Company.id == company_id).first()
    if not company:
        raise HTTPException(status_code=404, detail="Компания не найдена")
    if not company.website:
        raise HTTPException(status_code=400, detail="У компании нет сайта")
    
    try:
        from src.recon.web_crawler import crawl_website
        result = crawl_website(company.website, max_depth=2, max_pages=10)
        
        # Сохранить контакты
        for email in result.emails:
            db.add(Contact(company_id=company_id, type='email', value=email, source='web_crawl'))
        for phone in result.phones:
            db.add(Contact(company_id=company_id, type='phone', value=phone, source='web_crawl'))
        for platform, url in result.social_links.items():
            db.add(Contact(company_id=company_id, type=platform, value=url, label=platform.capitalize(), source='web_crawl'))
        
        if result.inn and not company.inn:
            company.inn = result.inn
        
        db.commit()
        
        return {
            "company": company.name,
            "emails": result.emails,
            "phones": result.phones,
            "social_links": result.social_links,
            "inn": result.inn
        }
    except Exception as e:
        return {"error": str(e)}


@app.get("/api/v1/recon/contacts")
def recon_contacts(db: Session = Depends(get_db), limit: int = 200):
    """Список всех найденных контактов с парсинга."""
    from sqlalchemy import text
    try:
        rows = db.execute(text("""
            SELECT c.value, c.type, c.source, c.label,
                   comp.name as company_name, comp.website
            FROM contacts c
            JOIN companies comp ON c.company_id = comp.id
            WHERE c.source = 'web_crawl'
            ORDER BY c.id DESC
            LIMIT :lim
        """), {"lim": limit}).fetchall()
        
        contacts = []
        for r in rows:
            contacts.append({
                "value": r[0],
                "type": r[1],
                "source": r[2],
                "label": r[3],
                "company_name": r[4],
                "website": r[5]
            })
        return {"contacts": contacts, "total": len(contacts)}
    except Exception as e:
        return {"error": str(e), "contacts": []}


# ─── Root redirect ───
@app.get("/")
def root():
    """Redirect to Mission Control dashboard."""
    from fastapi.responses import RedirectResponse
    return RedirectResponse(url="/app/")


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8001, reload=True)

